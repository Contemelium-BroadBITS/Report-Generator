from OpExpertOperations import Interactions

import os

from collections import Counter
from datetime import datetime
from inspect import currentframe
from subprocess import CalledProcessError, run
from yaml import safe_load

from docx2pdf import convert
from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import *
from docx.oxml.shared import OxmlElement, qn
from PyPDF2 import PdfReader
from spire.doc import *
from spire.doc.common import *

from pandas import DataFrame, to_datetime, to_numeric
from plotly.graph_objects import Figure, Pie, Scatter
from plotly.io import write_image
from plotly import express






class Console:
    
    def __init__(self):

        self.__initialize_local_variables()
        self.operator = Interactions()
        self.operator.login()


    
    def __initialize_local_variables(self):
        
        
        self.current_directory = os.path.dirname(os.path.abspath(__file__))
        
        self.current_name = str(datetime.now().strftime("%Y%m%d_%H%M%S"))
        
        if not os.path.exists(os.path.join(self.current_directory, 'templates')):
            os.makedirs(os.path.join(self.current_directory, 'templates'))
        self.template_directory = os.path.join(self.current_directory, 'templates')

        if not os.path.exists(os.path.join(self.current_directory, 'reports')):
            os.makedirs(os.path.join(self.current_directory, 'reports'))
        self.report_directory = os.path.join(self.current_directory, 'reports')
        
        

    def fetch_template(self, template_id):
        
        # template = DocxTemplate(template_path)
        if not os.path.exists(os.path.join(self.template_directory, f"{template_id}.docx")):
            downloaded_template = self.operator.downloadDocumentWithID(template_id)
            if downloaded_template.status_code == 200 and downloaded_template.content:
                with open(os.path.join(self.template_directory, f"{template_id}.docx"), "wb") as f:
                    f.write(downloaded_template.content)
                print(f"INFO at line {currentframe().f_lineno}: Template downloaded to {os.path.join(self.template_directory, f'{template_id}.docx')}")
                template_path = os.path.join(self.template_directory, f"{template_id}.docx")
                template = DocxTemplate(template_path)
            
        else:
            print(f"INFO at line {currentframe().f_lineno}: Template already exists at {os.path.join(self.template_directory, f'{template_id}.docx')}")
            template_path = os.path.join(self.template_directory, f"{template_id}.docx")
            template = DocxTemplate(template_path)
        
        return template


    def fetch_data_for_bar_chart(self, integration_id, abscissa = None, ordinate = None):

        data = self.operator.getIntegrationWithID(integration_id)
        if not data:
            print(f"ERROR at line {currentframe().f_lineno}: No data fetched with ID {integration_id}.")
            return None
        dataframe = DataFrame(data)
        columns = dataframe.columns.tolist()

        if abscissa and abscissa not in columns:
            print(f"Column '{abscissa}' does not exist in the dataset fetched with ID {integration_id}. Skipping...")
            return None
        if ordinate and ordinate not in columns:
            print(f"Column '{ordinate}' does not exist in the dataset fetched with ID {integration_id}. Skipping...")
            return None
        
        if not abscissa or not ordinate:
            if len(columns) < 2:
                print(f"ERROR at line {currentframe().f_lineno}: Not enough columns in the dataset fetched with ID {integration_id} to determine abscissa and ordinate.")
                return None
            abscissa = abscissa or columns[0]
            # Find the next column that is number-like
            ordinate = None
            for column in columns:
                if column == abscissa:
                    continue
                col_numeric = to_numeric(dataframe[column], errors='coerce').notna().all()
                if col_numeric:
                    ordinate = column
                    break
            if not ordinate:
                print(f"ERROR at line {currentframe().f_lineno}: No suitable numeric column found for ordinate in dataset fetched with ID {integration_id}.")
                return None

        # Ensure ordinate is numeric
        dataframe[ordinate] = to_numeric(dataframe[ordinate], errors="coerce")
        formatted_data = {
            "abscissa": dataframe[abscissa].dropna().tolist(),
            "ordinate": dataframe[ordinate].dropna().tolist(),
            "abscissa_label": abscissa,
            "ordinate_label": ordinate
        }

        if len(formatted_data['abscissa']) != len(formatted_data['ordinate']):
            print(f"ERROR at line {currentframe().f_lineno}: Mismatched lengths between abscissa and ordinate in dataset fetched with ID {integration_id}.")
            minimum_length = min(len(formatted_data['abscissa']), len(formatted_data['ordinate']))
            formatted_data['abscissa'] = formatted_data['abscissa'][:minimum_length]
            formatted_data['ordinate'] = formatted_data['ordinate'][:minimum_length]

        return formatted_data


    def fetch_data_for_stacked_bar_chart(self, integration_id, abscissa = None, ordinate = None):

        data = self.operator.getIntegrationWithID(integration_id)        

        return data


    def fetch_data_for_donut_chart(self, integration_id, column = None):
        
        data = self.operator.getIntegrationWithID(integration_id)
        if not data:
            print(f"ERROR at line {currentframe().f_lineno}: No data fetched with ID {integration_id}.")
            return None
        
        if column:            
            if column:
                if not any(column in item for item in data):
                    print(f"WARNING: Column '{column}' does not exist in the dataset fetched with ID {integration_id}. Skipping...")
                    return None            
            filtered_values = [item[column] for item in data if column in item]            
            counts = Counter(filtered_values)            
            formatted_data = {
                "labels": list(counts.keys()),
                "values": list(counts.values())
            }
            return formatted_data
        elif not column:
            formatted_data = {
                "labels": [],
                "values": []
            }
            for item in data:
                values = list(item.values())
                formatted_data["labels"].append(values[0])
                formatted_data["values"].append(values[1])
            return formatted_data


    def fetch_data_for_line_chart(self, integration_id, abscissa = None, ordinate = None):

        data = self.operator.getIntegrationWithID(integration_id)
        if not data:
            print(f"ERROR at line {currentframe().f_lineno}: No data fetched with ID {integration_id}.")
            return None
        dataframe = DataFrame(data)
        columns = dataframe.columns.tolist()
        
        if abscissa and abscissa not in columns:
            print(f"WARNING: Column '{abscissa}' does not exist in the dataset fetched with ID {integration_id}. Skipping...")
            return None
        if ordinate and ordinate not in columns:
            print(f"WARNING: Column '{ordinate}' does not exist in the dataset fetched with ID {integration_id}. Skipping...")
            return None
        
        if not abscissa or not ordinate:
            if len(columns) < 2:
                print(f"ERROR at line {currentframe().f_lineno}: Not enough columns in the dataset fetched with ID {integration_id} to determine abscissa and ordinate.")
                return None
            column_1 = abscissa or columns[0]
            column_2 = ordinate or columns[1]

        column_1_as_time = to_datetime(dataframe[column_1], errors = 'coerce').notna().all()
        column_2_as_time = to_datetime(dataframe[column_2], errors = 'coerce').notna().all()
        column_1_as_numeric = to_numeric(dataframe[column_1], errors = 'coerce').notna().all()
        column_2_as_numeric = to_numeric(dataframe[column_2], errors = 'coerce').notna().all()

        if column_1_as_time and not column_2_as_time:
            abscissa, ordinate = column_1, column_2
            dataframe[abscissa] = to_datetime(dataframe[abscissa], errors="coerce")
            dataframe[ordinate] = to_numeric(dataframe[ordinate], errors="coerce")
        elif column_2_as_time and not column_1_as_time:
            abscissa, ordinate = column_2, column_1
            dataframe[abscissa] = to_datetime(dataframe[abscissa], errors="coerce")
            dataframe[ordinate] = to_numeric(dataframe[ordinate], errors="coerce")
        elif column_1_as_numeric and column_2_as_numeric:
            abscissa, ordinate = column_1, column_2
            dataframe[abscissa] = to_numeric(dataframe[abscissa], errors="coerce")
            dataframe[ordinate] = to_numeric(dataframe[ordinate], errors="coerce")
        else:
            print(f"ERROR: Could not auto-detect suitable abscissa and ordinate.")
            return None
        try:
            dataframe = dataframe.sort_values(abscissa)
        except Exception as error:
            print(f"WARNING: Could not sort dataframe by abscissa '{abscissa}': {error}")

        # Format the data so that the entries are reduced to less than 20 points to have a spline like chart.
        if len(dataframe) > 20:
            dataframe = dataframe.iloc[::len(dataframe) // 20, :]

        formatted_data = {
            "abscissa": dataframe[abscissa].dropna().tolist(),
            "ordinate": dataframe[ordinate].dropna().tolist(), 
            "abscissa_label": abscissa,
            "ordinate_label": ordinate
        }
        
        
        # abscissa = formatted_data['abscissa']
        # has_timestamps = any(hasattr(value, "to_pydatetime") for value in abscissa)
        # if has_timestamps:
        #     normalized_abscissa = []
        #     for value in abscissa:
        #         try:
        #             if hasattr(value, "to_pydatetime"):
        #                 normalized_abscissa.append(value.to_pydatetime())
        #             else:
        #                 normalized_abscissa.append(value)
        #         except Exception:
        #             normalized_abscissa.append(value)
        #     formatted_data['abscissa'] = normalized_abscissa
        # else:
        #     try:
        #         formatted_data['abscissa'] = [float(value) for value in abscissa]
        #     except ValueError:
        #         pass
            
        if len(formatted_data['abscissa']) != len(formatted_data['ordinate']):
            print(f"ERROR at line {currentframe().f_lineno}: Mismatched lengths between abscissa and ordinate in dataset fetched with ID {integration_id}.")
            minimum_length = min(len(formatted_data['abscissa']), len(formatted_data['ordinate']))
            formatted_data['abscissa'] = formatted_data['abscissa'][:minimum_length]
            formatted_data['ordinate'] = formatted_data['ordinate'][:minimum_length]

        return formatted_data


    def fetch_data_for_pie_chart(self, integration_id, column = None):
        
        data = self.operator.getIntegrationWithID(integration_id)
        if not data:
            print(f"ERROR at line {currentframe().f_lineno}: No data fetched with ID {integration_id}.")
            return None
        
        if column:            
            if column:
                if not any(column in item for item in data):
                    print(f"WARNING: Column '{column}' does not exist in the dataset fetched with ID {integration_id}. Skipping...")
                    return None            
            filtered_values = [item[column] for item in data if column in item]            
            counts = Counter(filtered_values)            
            formatted_data = {
                "labels": list(counts.keys()),
                "values": list(counts.values())
            }
            return formatted_data
        elif not column:
            formatted_data = {
                "labels": [],
                "values": []
            }
            for item in data:
                values = list(item.values())
                formatted_data["labels"].append(values[0])
                formatted_data["values"].append(values[1])
            return formatted_data

    
    def fetch_data_for_table(self, integration_id, columns = []):
        
        data = self.operator.getIntegrationWithID(integration_id)
        
        if columns:
            filtered_data = []
            for item in data:
                filtered_item = {column: item[column] for column in columns if column in item}
                filtered_data.append(filtered_item)
            return filtered_data
        else:
            return data





class ReportGenerator(Console):
    
    def __init__(self):

        super().__init__()

        self.__initialize_local_variables()

    
    
    def __initialize_local_variables(self):
        
        self.supported_payloads = [
            'yaml-file', 
            'yaml-string'
        ]
        
        self.available_chart_types = [
            'pie',
            'donut',
            'bar',
            'line'
        ]
        
        self.available_palettes = {
            'crimson_red': {
                'colors': ['#DC143C', '#B22222', '#8B0000', '#FF6B6B', '#FF8E8E'],
                'background': '#FFFFFF'
            },
            'ocean_blue': {
                'colors': ['#0066CC', '#0080FF', '#3399FF', '#66B2FF', '#99CCFF'],
                'background': '#FFFFFF'
            },
            'violet': {
                'colors': ['#8A2BE2', '#9370DB', '#BA55D3', '#DA70D6', '#EE82EE'],
                'background': '#FFFFFF'
            },
            'emerald_green': {
                'colors': ['#50C878', '#4CAF50', '#388E3C', '#2E7D32', '#1B5E20'],
                'background': '#FFFFFF'
            },
            'sunset_orange': {
                'colors': ['#FF4500', '#FF6347', '#FF7F50', '#FF8C00', '#FFA07A'],
                'background': '#FFFFFF'
            }
        }
        
        default_palette = 'crimson_red'
        self.default_palette = {
            'default': default_palette,
            'pie': default_palette,
            'donut': default_palette,
            'bar': default_palette,
            'line': default_palette,
            'table': default_palette
        }
        
        self.default_legend = {
            'default': None,
            'pie': None,
            'donut': None,
            'bar': None,
            'line': None
        }
        
        self.rendered_document = None
        self.current_composition = None
    
    
    
    def __save_file_as_pdf(self, file_path, output_path = None):
        try:
            command = [
                'libreoffice', 
                '--headless', 
                '--convert-to', 'pdf', 
                file_path, 
                '--outdir', os.path.dirname(file_path)
            ]
            run(command, check = True)
        except CalledProcessError as e:
            print(f"ERROR at line {currentframe().f_lineno}: Failed to convert to PDF: {e}")
        except Exception as e:
            print(f"ERROR at line {currentframe().f_lineno}: An unexpected error occurred during PDF conversion: {e}")
            
    
    def __update_table_of_contents(self, file_path):
        ...
    
    
    def docx_to_pdf_pandoc(self, input_path, output_path):
        try:
            run(['pandoc', input_path, '-o', output_path], check=True)
            print(f"PDF saved to {output_path}")
        except Exception as e:
            print(f"Conversion failed: {e}")
            
    def docx_to_pdf_abiword(self, input_path, output_path):
        """
        Converts a DOCX file to PDF using AbiWord.
        """
        try:
            run([
                'abiword',
                '--to=pdf',
                input_path,
                '--to-name=' + output_path
            ], check=True)
            print(f"PDF saved to {output_path}")
        except Exception as e:
            print(f"Conversion failed: {e}")
            
    def docx_to_pdf_spire(self, input_path, output_path):
        
        document = Document()
        document.LoadFromFile(input_path)
        document.SaveToFile(output_path, FileFormat.PDF)   
        print(f"PDF saved to {output_path}")
        document.Close()

    def docx_to_pdf_docx2pdf(self, input_path, output_path):
        try:
            convert(input_path, output_path)
            print(f"PDF saved to {output_path}")
        except Exception as e:
            print(f"Conversion failed: {e}")

    def __determine_default_values(self, configuration):
        
        self.default_palette['default'] = configuration.get('defaultPalette', 'crimson_red')
        self.default_palette['pie'] = configuration.get('defaultPalette', 'crimson_red')
        self.default_palette['donut'] = configuration.get('defaultPalette', 'crimson_red')
        self.default_palette['bar'] = configuration.get('defaultPalette', 'crimson_red')
        self.default_palette['line'] = configuration.get('defaultPalette', 'crimson_red')
        
        self.default_legend['default'] = True
        self.default_legend['pie'] = True
        self.default_legend['donut'] = True
        self.default_legend['bar'] = True
        self.default_legend['line'] = True

        default_configuration = configuration.get('defaultConfiguration', [])
        for defaults in default_configuration:
            chart_type = defaults.get('type', '')
            if chart_type in self.available_chart_types:
                self.default_palette[chart_type] = defaults.get('palette', self.default_palette[chart_type])
                self.default_legend[chart_type] = defaults.get('legend', self.default_legend[chart_type])
    
    
    def __determine_palette_from_cell(self, cell):
        
        def hex_to_rgb(color):
            color = color.lstrip('#')
            return tuple(int(color[index:index + 2], 16) for index in (0, 2, 4))

        def rgb_distance(value_a, value_b):
            return sum((a - b) ** 2 for a, b in zip(value_a, value_b)) ** 0.5
        
        font_color = None
        try:
            # Get font color from the first run in the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    color = run.font.color
                    if color and color.rgb:
                        font_color = tuple(int(x) for x in color.rgb)
                        break
                if font_color:
                    break
        except Exception:
            pass

        if font_color is None:
            return self.default_palette.get('default', 'crimson_red')

        closest_palette = None
        min_distance = float('inf')
        for palette_name, palette in self.available_palettes.items():
            palette_rgb = hex_to_rgb(palette['colors'][0])
            dist = rgb_distance(font_color, palette_rgb)
            if dist < min_distance:
                min_distance = dist
                closest_palette = palette_name

        return closest_palette
               

    def __determine_palette_from_paragraph(self, paragraph):
        
        # Helper to convert hex color to RGB tuple
        def hex_to_rgb(color):
            color = color.lstrip('#')
            return tuple(int(color[index:index + 2], 16) for index in (0, 2, 4))
        
        # Helper to compute Euclidean distance between two RGB colors
        def rgb_distance(value_a, value_b):
            return sum((a - b) ** 2 for a, b in zip(value_a, value_b)) ** 0.5
        
        font_color = None
        try:
            # Get font color from the first run in the paragraph
            for run in paragraph.runs:
                color = run.font.color
                if color and color.rgb:
                    font_color = tuple(int(value) for value in color.rgb)
                    break
        except Exception:
            pass
        
        if font_color is None:
            return self.default_palette.get('table', 'crimson_red')
        
        closest_palette = None
        min_distance = float('inf')
        for palette_name, palette in self.available_palettes.items():
            palette_rgb = hex_to_rgb(palette['colors'][0])
            dist = rgb_distance(font_color, palette_rgb)
            if dist < min_distance:
                min_distance = dist
                closest_palette = palette_name
        return closest_palette
    
    
    def __locate_target_cell_and_size(self, template, key_variable, structure):
        
        target_cell = None
        current_table = None
        row_idx = col_idx = None
        try:
            tables = getattr(template.get_docx(), "tables", []) or []
        except Exception:
            tables = []

        placeholder = f"{key_variable}_image"
        for tbl in tables:
            found = False
            for r_i, row in enumerate(tbl.rows):
                for c_i, cell in enumerate(row.cells):
                    try:
                        if placeholder in cell.text:
                            target_cell = cell
                            current_table = tbl
                            row_idx, col_idx = r_i, c_i
                            found = True
                            break
                    except Exception:
                        continue
                if found:
                    break
            if found:
                break

        # width (inches) : prefer target_cell.width, else compute usable page width / cols
        width_in = None
        try:
            if target_cell is not None and getattr(target_cell, "width", None) is not None and getattr(target_cell.width, "inches", None) is not None:
                width_in = float(target_cell.width.inches)
        except Exception:
            width_in = None

        if width_in is None:
            usable = 6.0
            try:
                sections = getattr(template.get_docx(), "sections", []) or []
                if sections:
                    sec = sections[0]
                    usable = max(float(sec.page_width.inches) - float(sec.left_margin.inches) - float(sec.right_margin.inches), 1.0)
            except Exception:
                pass

            cols = 1
            try:
                if current_table is not None:
                    cols = len(current_table.columns) if hasattr(current_table, "columns") and len(current_table.columns) else len(current_table.rows[0].cells)
            except Exception:
                cols = 1
            width_in = (usable / max(cols, 1))
        # height (inches) : YAML override -> row.height -> XML fallbacks -> width fallback
        height_in = None
        height_source = None
        if isinstance(structure, dict):
            if "height_in" in structure:
                try:
                    height_in = float(structure.get("height_in"))
                    height_source = "override:height_in"
                except Exception:
                    height_in = None
            elif "height_px" in structure:
                try:
                    height_in = float(structure.get("height_px")) / 96.0
                    height_source = "override:height_px"
                except Exception:
                    height_in = None

        if height_in is None:
            try:
                if row_idx is not None and current_table is not None:
                    row = current_table.rows[row_idx]
                    if getattr(row, "height", None) is not None and getattr(row.height, "inches", None) is not None:
                        height_in = float(row.height.inches)
                        height_source = "template:row_height"
            except Exception:
                pass

        if height_in is None:
            # xml fallbacks (trHeight / tcW)
            import re
            try:
                if row_idx is not None and current_table is not None:
                    row_ref = current_table.rows[row_idx]
                    if hasattr(row_ref, "_tr"):
                        m = re.search(r'w:trHeight[^>]*w:val\s*=\s*"?(\d+)"?', row_ref._tr.xml or "")
                        if m:
                            height_in = int(m.group(1)) / 1440.0
                            height_source = "xml:trHeight"
            except Exception:
                pass

        if height_in is None:
            try:
                if target_cell is not None and hasattr(target_cell, "_tc"):
                    tc_xml = target_cell._tc.xml or ""
                    m2 = re.search(r'w:tcW[^>]*w:w\s*=\s*"?(\d+)"?', tc_xml)
                    if m2:
                        parsed = int(m2.group(1)) / 1440.0
                        if parsed > 0:
                            height_in = parsed
                            height_source = "xml:tcW_as_height"
            except Exception:
                pass

        if height_in is None:
            height_in = width_in
            height_source = "fallback:width"

        dpi = 96
        width_px = int(width_in * dpi)
        height_px = int(height_in * dpi)

        # prepare InlineImage sizes (docx.shared.Inches)
        from docx.shared import Inches
        try:
            inline_width = target_cell.width if getattr(target_cell, "width", None) is not None and getattr(target_cell.width, "inches", None) is not None else Inches(width_in)
        except Exception:
            inline_width = Inches(width_in)
        inline_height = Inches(height_in)

        # palette helper (call original routine if present, swallow errors)
        palette = None
        try:
            palette = self.__determine_palette_from_cell(target_cell)
        except Exception:
            palette = None

        return {
            "target_cell": target_cell,
            "table": current_table,
            "row": row_idx,
            "col": col_idx,
            "width_in": width_in,
            "height_in": height_in,
            "width_px": width_px,
            "height_px": height_px,
            "inline_width": inline_width,
            "inline_height": inline_height,
            "height_source": height_source,
            "palette": palette,
        }


    def __create_bar_chart(self, data, height, width, palette):
        
        if not data or not all(key in data for key in ['abscissa', 'ordinate', 'abscissa_label', 'ordinate_label']):
            print(f"ERROR at line {currentframe().f_lineno}: Invalid or incomplete data provided for bar chart.")
            return None
        
        palette = self.available_palettes.get(palette, self.available_palettes['crimson_red'])['colors']
        
        figure = express.bar(
            x = data["abscissa"], 
            y = data["ordinate"], 
            labels = {
                "x": data["abscissa_label"], 
                "y": data["ordinate_label"]
            },
            color_discrete_sequence = palette
        )
        figure.update_layout(
            height = height, 
            width = width, 
            showlegend = False,
            xaxis = dict(tickangle = -45)
        )
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]
        output_path = os.path.join(self.current_report_directory, f'{timestamp}.png')
        # figure.write_image(output_path)
        try:
            figure.write_image(output_path, width=width, height=height, scale=1)
        except Exception as e:
            try:
                write_image(figure, output_path, width=width, height=height, scale=1)
            except Exception as e2:
                print(f"ERROR at line {currentframe().f_lineno}: Failed to write bar chart image: {e} / {e2}")
                return None
        print(f"INFO at line {currentframe().f_lineno}: Bar chart saved to {output_path}")
        return output_path


    def __create_donut_chart(self, data, height, width, palette):
        """
        Generate a donut chart from formatted_data and save as PNG.

        Args:
            data (dict): The data for the donut chart.
            height (int): The height of the chart.
            width (int): The width of the chart.
            palette (list): The color palette to use.
        """
        if not data or not all(key in data for key in ['labels', 'values']):
            print(f"ERROR at line {currentframe().f_lineno}: Invalid or incomplete data provided for donut chart.")
            return None
        
        palette = self.available_palettes.get(palette, self.available_palettes['crimson_red'])['colors']
        
        figure = Figure(data = Pie(
            labels = data["labels"], 
            values = data["values"], 
            hole = 0.5,
            marker = dict(colors = palette), 
            textinfo = 'percent+label',
            insidetextorientation = 'auto', 
            domain = dict(x=[0, 1], y=[0, 1])
        ))
        figure.update_layout(
            height = height, 
            width = width, 
            margin = dict(l=4, r=4, t=4, b=4), 
            legend = dict(
                font = dict(size = 8),
                yanchor = "top",
                y = 0.99,
                xanchor = "left",
                x = 0.99
            ), 
            showlegend = False
        )
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]
        output_path = os.path.join(self.current_report_directory, f'{timestamp}.png')
        # figure.write_image(output_path)
        try:
            figure.write_image(output_path, width=width, height=height, scale=1)
        except Exception as e:
            try:
                write_image(figure, output_path, width=width, height=height, scale=1)
            except Exception as e2:
                print(f"ERROR at line {currentframe().f_lineno}: Failed to write donut chart image: {e} / {e2}")
                return None
        print(f"INFO at line {currentframe().f_lineno}: Donut chart saved to {output_path}")
        return output_path


    def __create_line_chart(self, data, height, width, palette):
        """
        Generate a line chart marking all points, showing value labels above each point,
        and circling the highest peak and lowest drop with palette shades.
        Value labels are staggered to reduce intersection.
        """
        if not data or not all(key in data for key in ['abscissa', 'ordinate', 'abscissa_label', 'ordinate_label']):
            print(f"ERROR at line {currentframe().f_lineno}: Invalid or incomplete data provided for line chart.")
            return None

        # Get palette colors
        try:
            if isinstance(palette, (list, tuple)):
                colors = list(palette)
            else:
                colors = self.available_palettes.get(palette, self.available_palettes['crimson_red'])['colors']
        except Exception:
            colors = self.available_palettes['crimson_red']['colors']

        x = data["abscissa"] or []
        y = data["ordinate"] or []
        if len(x) != len(y):
            minlen = min(len(x), len(y))
            x = x[:minlen]
            y = y[:minlen]
            print(f"WARNING at line {currentframe().f_lineno}: Trimmed abscissa/ordinate to same length ({minlen}).")

        # Assign marker colors from palette
        marker_colors = [colors[i % len(colors)] for i in range(len(x))]

        # Find maxima and minima indices
        if y:
            max_y = max(y)
            min_y = min(y)
            max_indices = [i for i, val in enumerate(y) if val == max_y]
            min_indices = [i for i, val in enumerate(y) if val == min_y]
        else:
            max_indices = []
            min_indices = []

        # Stagger text positions to reduce overlap
        text_positions = []
        for i in range(len(y)):
            if i in max_indices:
                text_positions.append("top right")
            elif i in min_indices:
                text_positions.append("bottom right")
            else:
                text_positions.append("top center" if i % 2 == 0 else "bottom center")

        # Main trace: all points, spline, value labels
        trace = Scatter(
            x=x,
            y=y,
            mode="lines+markers+text",
            line=dict(color=colors[0], width=2.4, shape="spline", smoothing=1.3),
            marker=dict(
                size=7,
                color=marker_colors,
                line=dict(width=0, color="black"),
            ),
            text=[f"{val:.2f}" for val in y],
            textposition=text_positions,
            textfont=dict(size=10, color="#222222"),
            hovertemplate=f"{data['abscissa_label']}: %{{x}}<br>{data['ordinate_label']}: %{{y}}<extra></extra>"
        )

        # Circle maxima (highest peak) with palette shade
        maxima_trace = Scatter(
            x=[x[i] for i in max_indices],
            y=[y[i] for i in max_indices],
            mode="markers",
            marker=dict(
                size=24,
                color=colors[1 % len(colors)] if len(colors) > 1 else "rgba(255,0,0,0.2)",
                line=dict(width=3, color=colors[2 % len(colors)] if len(colors) > 2 else "red"),
                symbol="circle-open"
            ),
            showlegend=False,
            hoverinfo="skip"
        )

        # Circle minima (lowest drop) with palette shade
        minima_trace = Scatter(
            x=[x[i] for i in min_indices],
            y=[y[i] for i in min_indices],
            mode="markers",
            marker=dict(
                size=24,
                color=colors[-1] if len(colors) > 1 else "rgba(0,0,255,0.2)",
                line=dict(width=3, color=colors[-2] if len(colors) > 2 else "blue"),
                symbol="circle-open"
            ),
            showlegend=False,
            hoverinfo="skip"
        )

        figure = Figure(data=[trace, maxima_trace, minima_trace])

        # Layout
        figure.update_layout(
            template = "simple_white",
            height = height,
            width = width,
            margin = dict(l = 20, r = 20, t = 20, b = 20),
            paper_bgcolor = "white",
            plot_bgcolor = "white",
            showlegend = False,
            font = dict(family="Inter, Helvetica, Arial, sans-serif", size=11, color="#222222"),
            hovermode = "x unified"
        )

        grid_color = "rgba(0,0,0,0.06)"
        figure.update_xaxes(
            title=dict(text=data["abscissa_label"], standoff=8),
            showgrid=True,
            gridcolor=grid_color,
            zeroline=False,
            showline=False,
            ticks="outside",
            ticklen=4,
            tickfont=dict(size=10, color="#444444"),
        )
        figure.update_yaxes(
            title=dict(text=data["ordinate_label"], standoff=8),
            showgrid=True,
            gridcolor=grid_color,
            zeroline=False,
            showline=False,
            ticks="outside",
            ticklen=0,
            tickfont=dict(size=10, color="#444444"),
        )

        # If abscissa is not numeric, treat as date
        try:
            import numbers
            is_numeric = all(isinstance(v, numbers.Number) for v in x) if x else False
            if not is_numeric:
                figure.update_xaxes(type="date", tickformatstops=None)
        except Exception:
            pass

        # Save image
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]
        output_path = os.path.join(self.current_report_directory, f'{timestamp}.png')
        try:
            figure.write_image(output_path, width=width, height=height, scale=1)
        except Exception as e:
            try:
                write_image(figure, output_path, width=width, height=height, scale=1)
            except Exception as e2:
                print(f"ERROR at line {currentframe().f_lineno}: Failed to write line chart image: {e} / {e2}")
                return None

        print(f"INFO at line {currentframe().f_lineno}: Line chart saved to {output_path}")
        return output_path
    
    
    def __create_pie_chart(self, data, height, width, palette):
        """
        Generate a pie chart from formatted_data and save as PNG.
        
        Args:
            data (dict): The data for the pie chart.
            height (int): The height of the chart.
            width (int): The width of the chart.
            palette (list): The color palette to use.
        """
        if not data or not all(key in data for key in ['labels', 'values']):
            print(f"ERROR at line {currentframe().f_lineno}: Invalid or incomplete data provided for pie chart.")
            return None
        
        palette = self.available_palettes.get(palette, self.available_palettes['crimson_red'])['colors']
        
        figure = Figure(data = Pie(
            labels = data["labels"], 
            values = data["values"], 
            marker = dict(colors = palette), 
            textinfo = 'percent+label',
            insidetextorientation = 'auto', 
            domain = dict(x=[0, 1], y=[0, 1])
        ))
        figure.update_layout(
            height = height, 
            width = width, 
            margin = dict(l=4, r=4, t=4, b=4),
            legend = dict(
                font = dict(size = 8),
                yanchor = "top",
                y = 0.99,
                xanchor = "left",
                x = 0.99
            ), 
            showlegend = False
        )
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]
        output_path = os.path.join(self.current_report_directory, f'{timestamp}.png')
        # figure.write_image(output_path)
        try:
            figure.write_image(output_path, width=width, height=height, scale=1)
        except Exception as e:
            try:
                write_image(figure, output_path, width=width, height=height, scale=1)
            except Exception as e2:
                print(f"ERROR at line {currentframe().f_lineno}: Failed to write pie chart image: {e} / {e2}")
                return None
        print(f"INFO at line {currentframe().f_lineno}: Pie chart saved to {output_path}")
        return output_path
    
    
    def __create_table(self, configuration, template):
        """
        Creates a table directly in the Word template with styling based on title color palette.
        """
        
        # Get data for the table
        data = self.fetch_data_for_table(configuration.get('data', ''), *configuration.get('columns', []))
        
        # Get specified columns from configuration, or use all available columns if not specified
        columns = configuration.get('columns', [])
        
        # If no columns specified, get all available columns from the first data entry
        if not columns and data:
            columns = list(data[0].keys())
        elif not columns:
            print("No data available for table")
            return
        
        # Get the document from template
        doc = template.get_docx()
        
        # Find the table title to determine palette
        palette = 'crimson_red'  # Default palette
        try:
            # Look for table_title in the document
            for paragraph in doc.paragraphs:
                if 'table_color' in paragraph.text or any('table_color' in run.text for run in paragraph.runs):
                    # Find the cell containing the title and determine palette
                    palette = self.__determine_palette_from_paragraph(paragraph)
                    break
        except Exception as e:
            print(f"Could not determine palette from title: {e}")
        
        print(f"Table Palette: {palette}")
        
        # Get colors from the determined palette
        palette_colors = self.available_palettes[palette]['colors']
        
        # Helper function to lighten colors for alternating rows
        def lighten_color(hex_color, factor=0.3):
            """Lighten a hex color by blending with white"""
            hex_color = hex_color.lstrip('#')
            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            # Blend with white
            lightened = tuple(int(c + (255 - c) * factor) for c in rgb)
            return f"{lightened[0]:02x}{lightened[1]:02x}{lightened[2]:02x}".upper()
        
        # Convert hex to RGB
        def hex_to_rgb(hex_color):
            hex_color = hex_color.lstrip('#')
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        
        # Prepare palette-based colors for the table
        header_color = palette_colors[0]  # Primary color for header
        header_rgb = hex_to_rgb(header_color)
        
        # Create variations for alternating rows using palette colors
        row_color_1 = lighten_color(palette_colors[0], 0.8)  # Very light version of primary color
        row_color_2 = lighten_color(palette_colors[1] if len(palette_colors) > 1 else palette_colors[0], 0.9)  # Even lighter version
        
        # Create a new table (remove existing tables if any)
        if doc.tables:
            # Clear existing tables
            for table in doc.tables:
                table._element.getparent().remove(table._element)
        
        # Add new table with proper dimensions
        table = doc.add_table(rows=len(data) + 1, cols=len(columns))
        table.allow_autofit = True
        # table.style = 'Table Grid'  # Use built-in table style
        
        def set_cell_padding(cell, top=100, bottom=100, left=100, right=100):
            """
            Set cell padding in twips (1/20 of a point).
            Example: 100 = 5 points.
            """
            tcPr = cell._tc.get_or_add_tcPr()
            cellMar = OxmlElement('w:tcMar')
            for side, value in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
                mar = OxmlElement(f'w:{side}')
                mar.set(qn('w:w'), str(value))
                mar.set(qn('w:type'), 'dxa')
                cellMar.append(mar)
            tcPr.append(cellMar)
        
        # Style the table
        def apply_table_styling(table):
            # Set table border using palette color
            tbl = table._tbl
            tblBorders = OxmlElement('w:tblBorders')
            
            # Use palette color for borders
            border_color = header_color.lstrip('#')
            
            # Define border style with palette color
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')  # Slightly thicker border
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), border_color)
                tblBorders.append(border)
            
            tbl.tblPr.append(tblBorders)
        
        def style_header_row(row, header_rgb):
            """Apply header styling to a row with palette-based color"""
            for cell in row.cells:
                # Set background color for header using palette color
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:val'), 'clear')
                shading_elm.set(qn('w:color'), 'auto')
                # Convert RGB to hex for the header background
                header_hex = f"{header_rgb[0]:02x}{header_rgb[1]:02x}{header_rgb[2]:02x}".upper()
                shading_elm.set(qn('w:fill'), header_hex)
                cell._tc.get_or_add_tcPr().append(shading_elm)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_padding(cell, top=120, bottom=30, left=140, right=140)
                
                # Set text color and formatting
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.keep_together = True
                    paragraph.paragraph_format.word_wrap = True
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                        run.font.bold = True
                        run.font.size = Inches(0.12)
                        # Below will enable text wrapping in header cells
                        run.font.word_wrap = True
                        run.font.keep_together = True
                        run.font.name = 'Arial'
        
        def style_data_row(row, is_even=False, row_color_1=None, row_color_2=None):
            """Apply data row styling with palette-based alternating colors"""
            bg_color = row_color_1 if is_even else row_color_2
            
            for cell in row.cells:
                # Set background color using palette-based colors
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:val'), 'clear')
                shading_elm.set(qn('w:color'), 'auto')
                shading_elm.set(qn('w:fill'), bg_color)
                cell._tc.get_or_add_tcPr().append(shading_elm)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_padding(cell, top=100, bottom=10, left=120, right=120)
                
                # Set text formatting with palette-based text color
                text_color = hex_to_rgb(palette_colors[2] if len(palette_colors) > 2 else palette_colors[0])
                # Darken the text color for better readability
                dark_text_color = tuple(int(c * 0.3) for c in text_color)  # Make it darker
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.keep_together = True
                    paragraph.paragraph_format.word_wrap = True
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(*dark_text_color)
                        run.font.size = Inches(0.10)
        
        # Apply overall table styling
        apply_table_styling(table)
        
        # Add and style header row
        header_row = table.rows[0]
        for col_idx, column in enumerate(columns):
            if col_idx < len(header_row.cells):
                header_row.cells[col_idx].text = column.title()
        
        style_header_row(header_row, header_rgb)
        
        # Add and style data rows
        for row_idx, entry in enumerate(data):
            data_row = table.rows[row_idx + 1]  # +1 to skip header
            for col_idx, column in enumerate(columns):
                if col_idx < len(data_row.cells):
                    cell_value = str(entry.get(column, ''))
                    data_row.cells[col_idx].text = cell_value
                    data_row.cells[col_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    data_row.cells[col_idx].paragraphs[0].paragraph_format.keep_together = True
                    data_row.cells[col_idx].paragraphs[0].paragraph_format.word_wrap = True
            
            # Apply alternating row styling with palette colors
            style_data_row(data_row, is_even=(row_idx % 2 == 0), row_color_1=row_color_1, row_color_2=row_color_2)
    
    
    def __create_table_2(self, structure, template):
        
        # Helper function to lighten colors for alternating rows
        def lighten_color(hex_color, factor = 0.3):
            hex_color = hex_color.lstrip('#')
            rgb = tuple(int(hex_color[index:index + 2], 16) for index in (0, 2, 4))
            lightened = tuple(int(c + (255 - c) * factor) for c in rgb)
            return f"{lightened[0]:02x}{lightened[1]:02x}{lightened[2]:02x}".upper()
        
        # Helper function to convert hex to RGB
        def hex_to_rgb(hex_color):
            hex_color = hex_color.lstrip('#')
            return tuple(int(hex_color[index:index + 2], 16) for index in (0, 2, 4))
        
        def apply_table_styling(table, border_color):
            table_structure = table._tbl
            table_borders = OxmlElement('w:tblBorders')
            border_color = border_color.lstrip('#')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), border_color)
                table_borders.append(border)
            table_structure.tblPr.append(table_borders)
        
        data = self.fetch_data_for_table(structure.get('data', ''), *structure.get('columns', []))
        document = template.get_docx()
        
        try:
            for paragraph in document.paragraphs:
                if 'table_color' in paragraph.text or any('table_color' in run.text for run in paragraph.runs):
                    palette = self.__determine_palette_from_paragraph(paragraph)
                    break
        except Exception as error:
            print(f"Could not determine palette from title: {error}")
            palette = self.default_palette.get('table', 'crimson_red')        
        palette = self.available_palettes.get(palette, self.available_palettes['crimson_red'])['colors']
        
        header_color = palette[0]
        header_rgb = hex_to_rgb(header_color)
        odd_row_color = lighten_color(palette[1] if len(palette) > 1 else palette[0], 0.9)
        even_row_color = lighten_color(palette[0], 0.7)
        
        # Create a new table (remove existing tables if any)
        if document.tables:
            for table in document.tables:
                table._element.getparent().remove(table._element)
        table = document.add_table(rows = len(data) + 1, cols = len(structure.get('columns', [])))


    def __load_configuration(self, configuration, type):

        if type.lower() not in self.supported_payloads:
            raise ValueError(f"Unsupported payload type: {type}")

        if type == 'yaml-string':
            self.report_structure = safe_load(configuration)

        elif type == 'yaml-file':
            with open(configuration, 'r') as file:
                self.report_structure = safe_load(file)
        
        else:
            raise ValueError(f"ERROR at line {currentframe().f_lineno}: Unsupported payload type: {type}")


    def __process_configuration(self):
        
        os.makedirs(os.path.join(self.report_directory, f"Report_generated_at_{self.current_name}"))
        self.current_report_directory = os.path.join(self.report_directory, f"Report_generated_at_{self.current_name}")
        
        self.report_meta = {}
        configuration = next((item for item in self.report_structure if item.get('type') in ['config', 'configuration']), {})
        self.report_meta['report_name'] = configuration.get('report_name', f"{datetime.now().strftime('Report_%Y-%m-%d_%H-%M-%S')}")
        self.report_meta['created_at'] = configuration.get('created_at', datetime.now().strftime('%Y.%m.%d'))
        self.report_meta['title'] = configuration.get('report_title', 'A title was not provided.')
        self.report_meta['description'] = configuration.get('report_description', 'A description was not provided.')
        self.report_meta['issued_for'] = configuration.get('issued_for', 'Unknown')

        self.__determine_default_values(configuration)

        template_type = configuration.get('template_type', 'dynamic')
        if template_type is None:
            raise ValueError(f"ERROR at line {currentframe().f_lineno}: Template type is not defined in the configuration.")
        elif template_type.lower() not in ['static', 'dynamic']:
            raise ValueError(f"ERROR at line {currentframe().f_lineno}: Unsupported template type: {template_type}")
        elif template_type.lower() == 'static':
            self.__process_static_configuration()
        elif template_type.lower() == 'dynamic':
            self.__process_dynamic_configuration()
        


    def __process_static_configuration(self):
        ...
    
    
    def __process_dynamic_configuration(self):
        
        configuration = next((item for item in self.report_structure if item.get('type') in ['config', 'configuration']), {})
        cover_page_template_id = configuration.get('cover_template_id', None)
        end_page_template_id = configuration.get('endcard_template_id', None)
        if cover_page_template_id is None or end_page_template_id is None:
            raise ValueError(f"ERROR at line {currentframe().f_lineno}: Cover page or end page template ID is not defined in the configuration.")

        # Processing the cover page of a dynamic report structure payload
        self.current_document = self.fetch_template(cover_page_template_id)
        if not self.current_document:
            raise ValueError(f"ERROR at line {currentframe().f_lineno}: Failed to load cover page template.")
        context = {
            'date': self.report_meta.get('created_at', ''),
            'title': self.report_meta.get('title', ''),
            'description': self.report_meta.get('description', ''),
            'issued_for': self.report_meta.get('issued_for', '')
        }
        self.current_document.render(context)
        self.current_document.add_page_break()
        self.current_composition = Composer(self.current_document.docx)
        self.list_of_pages = []
        self.list_of_contents = []
        self.current_page = 0
        
        page_configuration = []
        for step in self.report_structure:
            if step.get('type').lower() in ['config', 'configuration']:
                continue
            elif step.get('type').lower() in ['chart', 'table']:
                page_configuration.append(step)
            else:
                print(f"WARNING at line {currentframe().f_lineno}: Unsupported step type: {step.get('type')}. This step will be ignored.")
                continue
        
        for page in page_configuration:
            step_name = page.get('step', '<UNKNOWN_STEP>')
            template = self.fetch_template(page.get('template_id', ''))
            if not template:
                print(f"ERROR at line {currentframe().f_lineno}: Failed to load template for step {step_name}. This step will be skipped.")
            
            context = {}
            if page.get('type').lower() == 'chart':
                caption_index = 0
                for index, structure in enumerate(page.get('charts', [])):
                    if structure.get('type', None).lower() == 'caption':
                        caption_index += 1
                        key_variable = f"caption{caption_index}"
                        title = structure.get('title', '')
                        description = structure.get('description', '')
                        context[f"{key_variable}_title"] = title
                        context[f"{key_variable}_description"] = description
                        add_to_contents = structure.get('add_to_contents', True)
                        if add_to_contents:
                            content_meta = {
                                'title': title, 
                                'page': self.current_page + 1
                            }
                            self.list_of_contents.append(content_meta)
                        continue
                    key_variable = f"chart{(index + 1) - caption_index}"
                    title = structure.get('title', f'Chart {(index + 1) - caption_index}')
                    description = structure.get('description', '')
                    context[f"{key_variable}_title"] = title
                    context[f"{key_variable}_description"] = description
                    add_to_contents = structure.get('add_to_contents', True)
                    
                    criteria = self.__locate_target_cell_and_size(template, key_variable, structure)
                    # print(criteria)
                    if criteria is None:
                        print(f"ERROR at line {currentframe().f_lineno}: Could not locate target cell for {key_variable} in step {step_name}. This chart will be skipped.")
                        continue
                    else:
                        try:
                            palette = criteria['palette']
                            width = criteria['width_px']
                            height = criteria['height_px']
                            inline_width = criteria['inline_width']
                            inline_height = criteria['inline_height']
                            # inline_width = criteria['width_in']
                            # inline_height = criteria['height_in']
                        except Exception:
                            print(f"ERROR at line {currentframe().f_lineno}: Could not determine size or palette for {key_variable} in step {step_name}. This chart will be skipped.")
                            continue
                    
                    chart_type = structure.get('type', None).lower()
                    if chart_type not in self.available_chart_types:
                        print(f"WARNING at line {currentframe().f_lineno}: Unsupported chart type: {chart_type} in step {step_name}. This chart will be ignored.")
                        continue
                    if chart_type == 'bar':
                        abscissa = structure.get('abscissa', [])
                        ordinate = structure.get('ordinate', [])
                        data = self.fetch_data_for_bar_chart(structure.get('data', ''), abscissa, ordinate)
                        chart_image = self.__create_bar_chart(data, height, width, palette)
                    elif chart_type == 'donut':
                        column = structure.get('column', '')
                        data = self.fetch_data_for_donut_chart(structure.get('data', ''), column)
                        chart_image = self.__create_donut_chart(data, height, width, palette)
                    elif chart_type == 'line':
                        abscissa = structure.get('abscissa', [])
                        ordinate = structure.get('ordinate', [])
                        data = self.fetch_data_for_line_chart(structure.get('data', ''), abscissa, ordinate)
                        chart_image = self.__create_line_chart(data, height, width, palette)
                    elif chart_type == 'pie':
                        column = structure.get('column', '')
                        data = self.fetch_data_for_pie_chart(structure.get('data', ''), column)
                        chart_image = self.__create_pie_chart(data, height, width, palette)
                    else:
                        print(f"WARNING at line {currentframe().f_lineno}: Unsupported chart type: {chart_type} in step {step_name}. This chart will be ignored.")
                        continue
                    print(f"Width = {width}, Height = {height}, Inline Width = {inline_width}, Inline Height = {inline_height}")
                    context[f"{key_variable}_image"] = InlineImage(template, chart_image, width = inline_width, height = inline_height) if chart_image else None
                    if add_to_contents:
                        content_meta = {
                            'title': title, 
                            'page': self.current_page + 1
                        }
                        self.list_of_contents.append(content_meta)
                    # print(f"INFO at line {currentframe().f_lineno}: Processed {chart_type} chart for step {step_name}.")
                    # print(f"INFO at line {currentframe().f_lineno}: Fetched data for {chart_type} chart in step {step_name}.\nFetched Data: {data}\n\n")
            
            elif page.get('type').lower() == 'table':
                title = page.get('title', 'Table')
                description = page.get('description', '')
                context['table_title'] = title
                context['table_description'] = description
                add_to_contents = page.get('add_to_contents', True)
                self.__create_table(page, template)
                print(f"INFO at line {currentframe().f_lineno}: Processed table for step {step_name}.")
                if add_to_contents:
                    content_meta = {
                        'title': title, 
                        'page': self.current_page + 1
                    }
                    self.list_of_contents.append(content_meta)
            
            template.get_docx().add_page_break()
            template.render(context)
            # Current time as name suffix to avoid overwriting
            temporary_save_name_docx = os.path.join(self.current_report_directory, f"temporary_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]}.docx")
            temporary_save_name_pdf = temporary_save_name_docx.replace('.docx', '.pdf')
            temporary_context = {
                'date': self.report_meta.get('created_at', ''),
                'title': self.report_meta.get('title', ''),
                'description': self.report_meta.get('description', ''),
                'issued_for': self.report_meta.get('issued_for', '')
            }
            temporary_document = self.fetch_template(cover_page_template_id)
            temporary_document.render(temporary_context)
            temporary_document.add_page_break()
            temporary_composition = Composer(temporary_document)
            # template.save(temporary_save_name_docx)
            temporary_composition.append(template.docx)
            temporary_composition.doc.save(temporary_save_name_docx)
            self.__save_file_as_pdf(temporary_save_name_docx, temporary_save_name_pdf)
            temporary_reader = PdfReader(temporary_save_name_pdf)
            page_count_for_this_step = len(temporary_reader.pages) - 1
            # Now, delete the temporary saves
            try:
                os.remove(temporary_save_name_docx)
                os.remove(temporary_save_name_pdf)
            except Exception as e:
                print(f"WARNING at line {currentframe().f_lineno}: Could not delete temporary files: {e}")
            # self.current_composition.append(template.docx)
            self.list_of_pages.append(template.docx)
            # self.current_composition.append(template.docx)
            self.current_page += page_count_for_this_step
            print(f"INFO at line {currentframe().f_lineno}: Processed step {step_name} with {page_count_for_this_step} page(s).")
            # template.save(os.path.join(self.current_report_directory, f"intermediate_{str(step_name).replace(' ', '_')}.docx"))
        
        def __generate_table_of_contents(list_of_contents, template, table_of_contents_page_count = 0):
            
            table_of_contents_entries = []
            
            document = template.doc
            
            # 1. Adding the title and styling it
            table_of_contents_title = document.add_paragraph('Table of Contents')
            table_of_contents_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = table_of_contents_title.runs[0]
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(150, 50, 50)  # Dark gray text
            document.add_paragraph()  # Adding an empty paragraph for spacing
            
            # Creating the table for the table of contents and formatting it
            table_of_contents_table = document.add_table(rows = 1, cols = 2)
            table_of_contents_table.allow_autofit = True
            # section_width = Inches(5.5)
            # page_width = Inches(1.0)
            section = document.sections[0]
            page_width = section.page_width.inches - (section.left_margin.inches + section.right_margin.inches)
            page_column_width = Inches(1.0)
            section_column_width = Inches(page_width - page_column_width.inches)
            table_of_contents_table.columns[0].width = section_column_width
            table_of_contents_table.columns[1].width = page_column_width
            
            # Adding in the header row and styling it
            header_row = table_of_contents_table.rows[0].cells
            header_row[0].text = "SECTION"
            header_row[1].text = "PAGE"
            for index, cell in enumerate(header_row):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.keep_together = True
                    paragraph.paragraph_format.word_wrap = True
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(50, 50, 50)  # Dark gray text
                        run.font.bold = True
                        run.font.size = Pt(12)
                        run.font.word_wrap = True
                        run.font.keep_together = True
                        run.font.name = 'Arial'
            header_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Adding in the entries and styling them
            table_of_contents_table.add_row().cells
            for entry in list_of_contents:
                title = entry.get('title', '<N/A>')
                page = entry.get('page', '<N/A>')
                entry_row = table_of_contents_table.add_row().cells
                entry_row[0].text = str(title)
                entry_row[1].text = str(page + table_of_contents_page_count + 1)
                
                for paragraph in entry_row[0].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.keep_together = True
                    paragraph.paragraph_format.word_wrap = True
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(50, 50, 50)  # Dark gray text
                        run.font.size = Pt(12)
                        run.font.word_wrap = True
                        run.font.keep_together = True
                        run.font.name = 'Arial'
                
                for paragraph in entry_row[1].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.keep_together = True
                    paragraph.paragraph_format.word_wrap = True
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(150, 150, 150)  # Dark gray text
                        run.font.size = Pt(12)
                        run.font.word_wrap = True
                        run.font.keep_together = True
                        run.font.name = 'Arial'
        
        temporary_save_name_docx = os.path.join(self.current_report_directory, f"table_of_contents_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]}.docx")
        temporary_save_name_pdf = temporary_save_name_docx.replace('.docx', '.pdf')
        temporary_context = {
            'date': self.report_meta.get('created_at', ''),
            'title': self.report_meta.get('title', ''),
            'description': self.report_meta.get('description', ''),
            'issued_for': self.report_meta.get('issued_for', '')
        }        
        temporary_document = self.fetch_template(cover_page_template_id)
        temporary_document.render(temporary_context)
        temporary_document.add_page_break()
        temporary_composition = Composer(temporary_document)
        __generate_table_of_contents(self.list_of_contents, temporary_composition)
        temporary_composition.doc.save(temporary_save_name_docx)
        self.__save_file_as_pdf(temporary_save_name_docx, temporary_save_name_pdf)
        temporary_reader = PdfReader(temporary_save_name_pdf)
        page_count_for_this_step = len(temporary_reader.pages) - 1
        print(f"INFO at line {currentframe().f_lineno}: Table of contents generated with {page_count_for_this_step} page(s).")
        # Now, delete the temporary saves
        try:
            os.remove(temporary_save_name_docx)
            # os.remove(temporary_save_name_pdf)
        except Exception as e:
            print(f"WARNING at line {currentframe().f_lineno}: Could not delete temporary files: {e}")
        __generate_table_of_contents(self.list_of_contents, self.current_composition, page_count_for_this_step)
        self.current_composition.doc.add_page_break()
        for page in self.list_of_pages:
            self.current_composition.append(page)
        
        end_page_template = self.fetch_template(end_page_template_id)
        if not end_page_template:
            raise ValueError(f"ERROR at line {currentframe().f_lineno}: Failed to load end page template.")
        temporary_context = {
            'date': self.report_meta.get('created_at', ''),
            'title': self.report_meta.get('title', ''),
            'description': self.report_meta.get('description', ''),
            'issued_for': self.report_meta.get('issued_for', '')
        }
        end_page_template.render(temporary_context)
        # self.current_composition.doc.add_page_break()
        self.current_composition.append(end_page_template.docx)
        
        output_path = os.path.join(self.current_report_directory, f"{self.report_meta.get('report_name', f'{self.current_name}')}.docx")
        self.current_composition.doc.save(output_path)
        pdf_output_path = os.path.join(self.current_report_directory, f"{self.report_meta.get('report_name', f'{self.current_name}')}.pdf")
        # self.__update_table_of_contents(output_path)
        self.__save_file_as_pdf(output_path, self.current_report_directory)
        # self.docx_to_pdf_pandoc(output_path, pdf_output_path)
        # self.docx_to_pdf_abiword(output_path, pdf_output_path)
        # self.docx_to_pdf_spire(output_path, pdf_output_path)
        # self.docx_to_pdf_docx2pdf(output_path, pdf_output_path)
        print(f"INFO at line {currentframe().f_lineno}: Report saved to {output_path}")
        print(f"\n\n{self.list_of_contents}")
        
    
    
    
    def generate_report(self, configuration, type, email = None):

        self.__load_configuration(configuration, type)
        self.__process_configuration()




















if __name__ == "__main__":
    
    generator = ReportGenerator()
    yaml_file = os.path.join(generator.current_directory, '2025091600.yaml')
    generator.generate_report(yaml_file, 'yaml-file')